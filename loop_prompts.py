# make sure you pip install openai and python-docx

import openai
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.section import WD_SECTION


# Define the API key (Replace with your own API key)
api_key = input("Enter your OpenAI API key: ")

# Prompt the user for the required inputs
god_name = input("Enter the name of the god you want to design for (e.g., Zeus): ")
pantheon = input("Enter the mythological pantheon (e.g., Greek, Roman, Norse, etc.): ")
attack_type = input("Enter the attack type (Ranged or Melee): ")
damage_type = input("Enter the damage type (Physical or Magical): ")
god_class = input("Enter the class (Hunter, Mage, Assassin, Guardian, or Warrior): ")

# Load the document structure from the JSON file
with open('document_structure.json', 'r') as json_file:
    document_structure = json.load(json_file)

# Dictionary to store the generated content for each section and subsection
document_content = {}

# Define the updated system role message with additional context
system_message = f"You are a Game Designer for the video game SMITE as well as a Ph.D. in Mythological Studies from Harvard, charged with designing the god {god_name} for SMITE. " \
                 f"The god belongs to the {pantheon} pantheon, has an {attack_type} attack type, " \
                 f"{damage_type} damage type, and is of the {god_class} class. Answer questions thoroughly. Do not use the first person. Write your answers as if intended to be be used in a portion of a design document."

# Create a new Word document

# Create a new Word document
doc = Document()

# Create title page
doc.add_section(WD_SECTION.NEW_PAGE)
title = doc.add_paragraph(f'Background and Brainstorm Document for {god_name}')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].font.size = Pt(24)
doc.add_page_break()

# Add table of contents
doc.add_paragraph('Table of Contents', style='Heading1')
doc.add_paragraph('<<coming soon>>')  # Placeholder for the table of contents
doc.add_page_break()

# Iterate through the document structure and call the ChatGPT API
for section_item in document_structure:
    section = section_item["section"]
    document_content[section] = {}
    doc.add_heading(section, level=2)
    for subsection_item in section_item["subsections"]:
        subsection = subsection_item["subsection"]
        # Insert the god's name into the prompt using string formatting
        user_prompt = subsection_item["prompt"].format(god_name=god_name)
        # Indicate the current section and subsection being processed
        print(f"Processing: Section '{section}', Subsection '{subsection}'")
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            api_key=api_key
        )
        # Store the generated content in the nested dictionary
        content = response['choices'][0]['message']['content']
        document_content[section][subsection] = content
        # Add the subsection heading and content to the Word document
        doc.add_heading(subsection, level=3)
        doc.add_paragraph(content)
    # Add a page break after each major section
    doc.add_page_break()

# Add the table of contents to the document
#doc.tables_of_contents[0].add_entries(doc.paragraphs, 2)

# Add Appendix heading
doc.add_heading('Appendix', level=1)

# Add JSON Data heading in the Appendix
doc.add_heading('JSON Data Used in Document Creation', level=2)

# Convert the JSON data (document_structure) to a formatted string
json_data_str = json.dumps(document_structure, indent=4)

# Add JSON data as preformatted text to the Word document
json_paragraph = doc.add_paragraph(json_data_str)
json_paragraph.style = 'NoSpacing'
json_paragraph.runs[0].font.name = 'Courier New'  # Monospace font for JSON data

# Save the Word document
doc.save(f'{god_name}_Design_Document.docx')

print("Design document has been successfully created and saved!")